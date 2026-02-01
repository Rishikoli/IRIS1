
import os
import json
import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

# Configuration
PROJECT_ROOT = "/home/aditya/IRIS1/frontend"
OUTPUT_FILE = "/home/aditya/.gemini/antigravity/brain/7dd913dd-55f5-469f-a71e-74da9c1a5644/frontend_tech_report.docx"

def get_dependencies():
    with open(os.path.join(PROJECT_ROOT, "package.json"), 'r') as f:
        data = json.load(f)
    return data.get('dependencies', {}), data.get('devDependencies', {})

def create_document():
    document = Document()
    
    # Title
    title = document.add_heading('IRIS1 Frontend Technical Architecture', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Metadata
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Generated: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    run.font.color.rgb = RGBColor(128, 128, 128)
    run.italic = True
    
    document.add_heading('1. Executive Summary', level=1)
    p = document.add_paragraph(
        'The IRIS1 frontend is a high-performance, interactive financial forensics dashboard built with '
    )
    p.add_run('Next.js 15').bold = True
    p.add_run(' and ')
    p.add_run('React 19').bold = True
    p.add_run('. It leverages advanced WebGL visualizations (')
    p.add_run('Three.js, React Sphere').italic = True
    p.add_run(') and network graph analysis (')
    p.add_run('React Flow, ShellHunter3D').italic = True
    p.add_run(') to detect financial anomalies and fraud rings in real-time. The UI follows a modern "Neumorphic" design language implemented with ')
    p.add_run('Tailwind CSS v4').bold = True
    p.add_run('.')

    document.add_heading('2. Architecture & Data Flow', level=1)
    
    document.add_heading('Core Stack', level=2)
    p = document.add_paragraph(style='List Bullet')
    p.add_run('Runtime: ').bold = True
    p.add_run('Node.js / Browser (Universal Rendering)')
    
    p = document.add_paragraph(style='List Bullet')
    p.add_run('Routing: ').bold = True
    p.add_run('Next.js App Router (File-system based)')
    
    p = document.add_paragraph(style='List Bullet')
    p.add_run('State Management: ').bold = True
    p.add_run('React Context + Local State')
    
    p = document.add_paragraph(style='List Bullet')
    p.add_run('Styling: ').bold = True
    p.add_run('Tailwind CSS v4 + PostCSS')

    document.add_heading('3. Component Deep Dives', level=1)
    
    # Liquid Ether
    document.add_heading('Liquid Ether Simulation (LiquidEther.tsx)', level=2)
    document.add_paragraph('A custom WebGL fluid dynamics engine used for the landing page background.')
    p = document.add_paragraph(style='List Bullet')
    p.add_run('Technology: ').bold = True
    p.add_run('Three.js, Custom GLSL Shaders')
    p = document.add_paragraph(style='List Bullet')
    p.add_run('Technique: ').bold = True
    p.add_run('FBO (Frame Buffer Object) ping-ponging for physics')
    
    # ShellHunter
    document.add_heading('Shell Hunter 3D (ShellHunter3D.tsx)', level=2)
    document.add_paragraph('Force-directed 3D graph analysis tool for detecting circular trading loops.')
    p = document.add_paragraph(style='List Bullet')
    p.add_run('Technology: ').bold = True
    p.add_run('react-force-graph-3d, Three.js')
    p = document.add_paragraph(style='List Bullet')
    p.add_run('Features: ').bold = True
    p.add_run('Cycle detection, 3D camera controls, Particle effects')

    document.add_heading('4. Comprehensive Dependency Matrix', level=1)
    deps, dev_deps = get_dependencies()
    
    # Prod Deps
    document.add_heading('Production Dependencies', level=2)
    table = document.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Package'
    hdr_cells[1].text = 'Version'
    hdr_cells[2].text = 'Purpose'
    
    for name, version in sorted(deps.items()):
        purpose = "Utility"
        if "react" in name: purpose = "UI Library"
        if "three" in name or "d3" in name or "chart" in name or "graph" in name: purpose = "Visualization"
        if "next" in name: purpose = "Framework"
        if "axios" in name: purpose = "Networking"
        
        row_cells = table.add_row().cells
        row_cells[0].text = name
        row_cells[1].text = version
        row_cells[2].text = purpose

    # Dev Deps
    document.add_heading('Development Dependencies', level=2)
    table = document.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Package'
    hdr_cells[1].text = 'Version'
    hdr_cells[2].text = 'Purpose'
    
    for name, version in sorted(dev_deps.items()):
        row_cells = table.add_row().cells
        row_cells[0].text = name
        row_cells[1].text = version
        row_cells[2].text = "Dev Tool"

    document.save(OUTPUT_FILE)
    print(f"DOCX Report generated at: {OUTPUT_FILE}")

if __name__ == "__main__":
    create_document()
